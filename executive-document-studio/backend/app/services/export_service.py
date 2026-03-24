from __future__ import annotations

import io
from html import escape

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models.schemas import Draft


class ExportService:
    def to_markdown_bytes(self, draft: Draft) -> bytes:
        return self._render_markdown(draft).encode('utf-8')

    def to_text_bytes(self, draft: Draft) -> bytes:
        lines = [draft.title, '=' * len(draft.title), '']
        for section, content in draft.content.items():
            pretty = section.replace('_', ' ').title()
            lines.append(pretty)
            lines.append('-' * len(pretty))
            lines.append(content)
            lines.append('')
        lines.append('Sources Used')
        lines.append('------------')
        for citation in draft.citations:
            lines.append(f'* {citation.document_title or citation.document_id}: {citation.snippet or citation.chunk_id}')
        return '\n'.join(lines).encode('utf-8')

    def to_html_bytes(self, draft: Draft) -> bytes:
        html_parts = [
            '<html><head><meta charset="utf-8"><title>{}</title></head><body style="font-family: Inter, Arial, sans-serif; padding: 32px;">'.format(self._escape(draft.title)),
            f'<h1>{self._escape(draft.title)}</h1>',
        ]
        for section, content in draft.content.items():
            html_parts.append(f'<h2>{self._escape(section.replace("_", " ").title())}</h2>')
            html_parts.append(self._paragraphs_to_html(content))
        html_parts.append('<h2>Sources Used</h2><ul>')
        for citation in draft.citations:
            html_parts.append(
                f'<li><strong>{self._escape(citation.document_title or citation.document_id)}</strong>: {self._escape(citation.snippet or citation.chunk_id)}</li>'
            )
        html_parts.append('</ul></body></html>')
        return ''.join(html_parts).encode('utf-8')

    def to_docx_bytes(self, draft: Draft) -> bytes:
        document = DocxDocument()
        document.add_heading(draft.title, level=0)
        for section, content in draft.content.items():
            document.add_heading(section.replace('_', ' ').title(), level=1)
            for paragraph in self._split_paragraphs(content):
                document.add_paragraph(paragraph)
        document.add_heading('Sources Used', level=1)
        for citation in draft.citations:
            document.add_paragraph(
                f'{citation.document_title or citation.document_id}: {citation.snippet or citation.chunk_id}',
                style='List Bullet',
            )
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def to_pdf_bytes(self, draft: Draft) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER, leftMargin=48, rightMargin=48, topMargin=48, bottomMargin=48)
        styles = getSampleStyleSheet()
        story = [Paragraph(self._escape(draft.title), styles['Title']), Spacer(1, 12)]
        for section, content in draft.content.items():
            story.append(Paragraph(self._escape(section.replace('_', ' ').title()), styles['Heading2']))
            story.append(Spacer(1, 6))
            for paragraph in self._split_paragraphs(content):
                story.append(Paragraph(self._escape(paragraph).replace('\n', '<br/>'), styles['BodyText']))
                story.append(Spacer(1, 8))
        story.append(Paragraph('Sources Used', styles['Heading2']))
        story.append(Spacer(1, 6))
        for citation in draft.citations:
            story.append(Paragraph(self._escape(f'{citation.document_title or citation.document_id}: {citation.snippet or citation.chunk_id}'), styles['BodyText']))
            story.append(Spacer(1, 4))
        doc.build(story)
        return buffer.getvalue()

    def _render_markdown(self, draft: Draft) -> str:
        lines = [f'# {draft.title}', '']
        for section, content in draft.content.items():
            pretty = section.replace('_', ' ').title()
            lines.append(f'## {pretty}')
            lines.append(content)
            lines.append('')
        lines.append('## Sources Used')
        for citation in draft.citations:
            lines.append(f'- {citation.document_title or citation.document_id}: {citation.snippet or citation.chunk_id}')
        return '\n'.join(lines)

    def _split_paragraphs(self, value: str) -> list[str]:
        paragraphs = [segment.strip() for segment in value.split('\n\n') if segment.strip()]
        return paragraphs or [value.strip()]

    def _paragraphs_to_html(self, value: str) -> str:
        paragraphs = self._split_paragraphs(value)
        return ''.join(f'<p>{self._escape(paragraph).replace(chr(10), "<br/>")}</p>' for paragraph in paragraphs)

    def _escape(self, value: str) -> str:
        return escape(value, quote=True)


export_service = ExportService()
