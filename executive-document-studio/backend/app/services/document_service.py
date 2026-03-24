from __future__ import annotations

import re
import tempfile
from datetime import datetime
from html import unescape
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, List
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import UploadFile
from PyPDF2 import PdfReader

from app.models.schemas import Chunk, Document, DocumentCreateText, DocumentCreateUrl, DocumentType, SearchResponse, SearchResult
from app.services import database
from app.services.audit_service import audit_service
from app.services.embedding_service import embedding_service


class ReadableHtmlExtractor(HTMLParser):
    BLOCK_TAGS = {
        'article', 'aside', 'blockquote', 'br', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'header', 'footer', 'li', 'main', 'nav', 'p', 'section', 'tr', 'ul', 'ol',
    }
    SKIP_TAGS = {'script', 'style', 'noscript', 'svg'}

    def __init__(self) -> None:
        super().__init__()
        self.text_parts: list[str] = []
        self.title_parts: list[str] = []
        self.skip_depth = 0
        self.in_title = False

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self.skip_depth += 1
            return
        if tag == 'title':
            self.in_title = True
            return
        if self.skip_depth == 0 and tag in self.BLOCK_TAGS:
            self.text_parts.append('\n')

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS and self.skip_depth > 0:
            self.skip_depth -= 1
            return
        if tag == 'title':
            self.in_title = False
            return
        if self.skip_depth == 0 and tag in self.BLOCK_TAGS:
            self.text_parts.append('\n')

    def handle_data(self, data: str) -> None:
        if self.skip_depth > 0:
            return
        cleaned = data.strip()
        if not cleaned:
            return
        if self.in_title:
            self.title_parts.append(cleaned)
        else:
            self.text_parts.append(cleaned + ' ')

    def extract(self) -> tuple[str, str]:
        title = unescape(' '.join(self.title_parts)).strip()
        text = unescape(''.join(self.text_parts))
        text = re.sub(r'\r', '', text)
        text = re.sub(r'[ \t]+\n', '\n', text)
        text = re.sub(r'\n[ \t]+', '\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return title, text.strip()


class DocumentService:
    HYBRID_KEYWORD_WEIGHT = 0.58
    HYBRID_VECTOR_WEIGHT = 0.42
    VECTOR_FALLBACK_THRESHOLD = 0.16

    async def list_documents(self) -> List[Document]:
        rows = await database.fetch_all('SELECT * FROM documents ORDER BY date DESC, title ASC')
        return [self._row_to_document(row) for row in rows]

    async def get_document(self, document_id: str) -> Document | None:
        row = await database.fetch_one('SELECT * FROM documents WHERE id = ?', (document_id,))
        return self._row_to_document(row) if row else None

    async def get_chunks_for_documents(self, document_ids: list[str]) -> list[Chunk]:
        if not document_ids:
            return []
        placeholders = ', '.join('?' for _ in document_ids)
        rows = await database.fetch_all(
            f'SELECT * FROM chunks WHERE document_id IN ({placeholders}) ORDER BY document_id, position ASC, id ASC',
            tuple(document_ids),
        )
        return [self._row_to_chunk(row) for row in rows]

    async def get_all_chunks(self) -> list[Chunk]:
        rows = await database.fetch_all('SELECT * FROM chunks ORDER BY document_id, position ASC, id ASC')
        return [self._row_to_chunk(row) for row in rows]

    async def get_chunk(self, chunk_id: str) -> Chunk | None:
        row = await database.fetch_one('SELECT * FROM chunks WHERE id = ?', (chunk_id,))
        return self._row_to_chunk(row) if row else None

    async def keyword_search(self, query: str, limit: int = 10, doc_ids: list[str] | None = None) -> list[SearchResult]:
        query_tokens = self._tokenize(query)
        chunks = await (self.get_chunks_for_documents(doc_ids) if doc_ids else self.get_all_chunks())

        scored: list[SearchResult] = []
        for chunk in chunks:
            score, matched_terms = self._score_chunk(query_tokens, chunk.text)
            if score <= 0:
                continue
            scored.append(
                SearchResult(
                    chunk_id=chunk.id,
                    document_id=chunk.document_id,
                    document_title=chunk.document_title,
                    text=chunk.text,
                    score=round(score, 4),
                    keyword_score=round(score, 4),
                    vector_score=0.0,
                    matched_terms=matched_terms,
                    metadata={**chunk.metadata, 'position': chunk.position, 'embedding_available': chunk.embedding is not None},
                )
            )
        return sorted(scored, key=lambda item: (item.score, len(item.text)), reverse=True)[:limit]

    async def search(self, query: str, limit: int = 10, doc_ids: list[str] | None = None) -> SearchResponse:
        chunks = await (self.get_chunks_for_documents(doc_ids) if doc_ids else self.get_all_chunks())
        query_tokens = self._tokenize(query)
        query_embedding = embedding_service.embed_text(query) if query_tokens else None

        staged: list[dict[str, Any]] = []
        max_keyword_raw = 0.0
        for chunk in chunks:
            keyword_raw, matched_terms = self._score_chunk(query_tokens, chunk.text)
            max_keyword_raw = max(max_keyword_raw, keyword_raw)
            vector_score = 0.0
            if query_embedding is not None and chunk.embedding:
                vector_score = max(0.0, embedding_service.cosine_similarity(query_embedding, chunk.embedding))
            staged.append(
                {
                    'chunk': chunk,
                    'keyword_raw': keyword_raw,
                    'matched_terms': matched_terms,
                    'vector_score': vector_score,
                }
            )

        results: list[SearchResult] = []
        for item in staged:
            chunk = item['chunk']
            keyword_score = (item['keyword_raw'] / max_keyword_raw) if max_keyword_raw > 0 and item['keyword_raw'] > 0 else 0.0
            vector_score = item['vector_score']
            hybrid_score = (keyword_score * self.HYBRID_KEYWORD_WEIGHT) + (vector_score * self.HYBRID_VECTOR_WEIGHT)
            if item['keyword_raw'] > 0:
                hybrid_score += 0.03

            if item['keyword_raw'] <= 0 and vector_score < self.VECTOR_FALLBACK_THRESHOLD:
                continue

            results.append(
                SearchResult(
                    chunk_id=chunk.id,
                    document_id=chunk.document_id,
                    document_title=chunk.document_title,
                    text=chunk.text,
                    score=round(min(hybrid_score, 1.0), 4),
                    keyword_score=round(keyword_score, 4),
                    vector_score=round(vector_score, 4),
                    matched_terms=item['matched_terms'],
                    metadata={
                        **chunk.metadata,
                        'position': chunk.position,
                        'embedding_available': chunk.embedding is not None,
                    },
                )
            )

        if not results and staged:
            fallback = sorted(
                staged,
                key=lambda item: (item['vector_score'], item['keyword_raw'], len(item['chunk'].text)),
                reverse=True,
            )[: min(limit, 5)]
            for item in fallback:
                chunk = item['chunk']
                keyword_score = (item['keyword_raw'] / max_keyword_raw) if max_keyword_raw > 0 and item['keyword_raw'] > 0 else 0.0
                hybrid_score = (keyword_score * self.HYBRID_KEYWORD_WEIGHT) + (item['vector_score'] * self.HYBRID_VECTOR_WEIGHT)
                results.append(
                    SearchResult(
                        chunk_id=chunk.id,
                        document_id=chunk.document_id,
                        document_title=chunk.document_title,
                        text=chunk.text,
                        score=round(min(hybrid_score, 1.0), 4),
                        keyword_score=round(keyword_score, 4),
                        vector_score=round(item['vector_score'], 4),
                        matched_terms=item['matched_terms'],
                        metadata={
                            **chunk.metadata,
                            'position': chunk.position,
                            'embedding_available': chunk.embedding is not None,
                            'fallback': True,
                        },
                    )
                )

        results = sorted(
            results,
            key=lambda item: (item.score, item.vector_score, item.keyword_score, len(item.text)),
            reverse=True,
        )[:limit]

        documents_hit = len({item.document_id for item in results})
        average_score = round(sum(item.score for item in results) / len(results), 4) if results else 0.0
        average_keyword_score = round(sum(item.keyword_score for item in results) / len(results), 4) if results else 0.0
        average_vector_score = round(sum(item.vector_score for item in results) / len(results), 4) if results else 0.0

        return SearchResponse(
            query=query,
            summary={
                'mode': 'hybrid',
                'keyword_weight': self.HYBRID_KEYWORD_WEIGHT,
                'vector_weight': self.HYBRID_VECTOR_WEIGHT,
                'vector_threshold': self.VECTOR_FALLBACK_THRESHOLD,
                'embedding_dimension': embedding_service.dimension,
                'candidates_scanned': len(chunks),
                'documents_filtered': len(doc_ids or []),
                'documents_hit': documents_hit,
                'returned': len(results),
                'average_score': average_score,
                'average_keyword_score': average_keyword_score,
                'average_vector_score': average_vector_score,
            },
            results=results,
        )

    async def create_text_document(self, payload: DocumentCreateText) -> Document:
        document_id = f'doc_{uuid4().hex[:10]}'
        now = datetime.utcnow().isoformat()
        summary = payload.summary or self._summarize(payload.content)
        await database.execute(
            '''
            INSERT INTO documents (id, title, type, tags, date, summary, content, source_name, source_kind, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                document_id,
                payload.title,
                payload.type.value,
                database.dumps(payload.tags),
                payload.date,
                summary,
                payload.content,
                payload.title,
                'text',
                now,
                now,
            ),
        )
        await self._replace_chunks(document_id, payload.title, payload.content)
        await embedding_service.index_document(document_id)
        await audit_service.log('document.created', 'document', document_id, {'source_kind': 'text'})
        document = await self.get_document(document_id)
        assert document is not None
        return document

    async def create_url_document(self, payload: DocumentCreateUrl) -> Document:
        fetched_title, content, final_url = self._fetch_url_content(payload.url)
        if not content.strip():
            raise ValueError('Unable to extract text from the provided URL')

        document_id = f'doc_{uuid4().hex[:10]}'
        now = datetime.utcnow().isoformat()
        title = payload.title.strip() if payload.title and payload.title.strip() else fetched_title
        summary = self._summarize(content)
        await database.execute(
            '''
            INSERT INTO documents (id, title, type, tags, date, summary, content, source_name, source_kind, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                document_id,
                title,
                payload.type.value,
                database.dumps(payload.tags),
                payload.date,
                summary,
                content,
                final_url,
                'url',
                now,
                now,
            ),
        )
        await self._replace_chunks(document_id, title, content)
        await embedding_service.index_document(document_id)
        await audit_service.log('document.ingested_url', 'document', document_id, {'url': final_url, 'source_kind': 'url'})
        document = await self.get_document(document_id)
        assert document is not None
        return document

    async def upload_document(self, file: UploadFile, document_type: DocumentType, tags: list[str], date: str | None = None) -> Document:
        if not file.filename:
            raise ValueError('Uploaded file must include a filename')

        suffix = Path(file.filename).suffix.lower()
        upload_path = database.UPLOADS_DIR / f'{uuid4().hex}{suffix}'
        content_bytes = await file.read()
        upload_path.write_bytes(content_bytes)
        content = self._extract_text(upload_path, suffix)
        if not content.strip():
            raise ValueError('Unable to extract any text from uploaded file')

        document_id = f'doc_{uuid4().hex[:10]}'
        now = datetime.utcnow().isoformat()
        title = Path(file.filename).stem.replace('_', ' ').replace('-', ' ').title()
        summary = self._summarize(content)
        await database.execute(
            '''
            INSERT INTO documents (id, title, type, tags, date, summary, content, source_name, source_kind, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''',
            (
                document_id,
                title,
                document_type.value,
                database.dumps(tags),
                date,
                summary,
                content,
                file.filename,
                suffix.lstrip('.') or 'file',
                now,
                now,
            ),
        )
        await self._replace_chunks(document_id, title, content)
        await embedding_service.index_document(document_id)
        await audit_service.log(
            'document.uploaded',
            'document',
            document_id,
            {'filename': file.filename, 'source_kind': suffix.lstrip('.') or 'file'},
        )
        document = await self.get_document(document_id)
        assert document is not None
        return document

    async def list_audit_logs(self, limit: int = 50):
        return await audit_service.list_recent(limit=limit)

    async def _replace_chunks(self, document_id: str, title: str, content: str) -> None:
        await database.execute('DELETE FROM chunks WHERE document_id = ?', (document_id,))
        chunks = self._chunk_text(content)
        rows = []
        for position, text in enumerate(chunks, start=1):
            rows.append(
                (
                    f'{document_id}_chunk_{position:03d}',
                    document_id,
                    title,
                    text,
                    position,
                    database.dumps({'position': position, 'word_count': len(text.split())}),
                    None,
                )
            )
        await database.insert_many(
            'INSERT INTO chunks (id, document_id, document_title, text, position, metadata, embedding) VALUES (?, ?, ?, ?, ?, ?, ?)',
            rows,
        )

    def _fetch_url_content(self, url: str) -> tuple[str, str, str]:
        parsed = urlparse(url)
        if parsed.scheme not in {'http', 'https'}:
            raise ValueError('Only http and https URLs are supported')

        request = Request(
            url,
            headers={
                'User-Agent': 'Mozilla/5.0 (compatible; ExecutiveDocumentStudio/1.0; +https://shakudo.io)',
                'Accept': 'text/html,application/pdf,text/plain,application/vnd.openxmlformats-officedocument.wordprocessingml.document,*/*',
            },
        )

        try:
            with urlopen(request, timeout=20) as response:
                final_url = response.geturl()
                content_type = (response.headers.get('Content-Type') or '').lower()
                payload = response.read()
        except HTTPError as exc:
            raise ValueError(f'Unable to fetch URL (HTTP {exc.code})') from exc
        except URLError as exc:
            raise ValueError(f'Unable to fetch URL: {exc.reason}') from exc

        suffix = Path(urlparse(final_url).path).suffix.lower()
        if 'application/pdf' in content_type or suffix == '.pdf':
            return self._extract_remote_file(final_url, payload, '.pdf')
        if 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type or suffix == '.docx':
            return self._extract_remote_file(final_url, payload, '.docx')
        if 'text/plain' in content_type or suffix == '.txt':
            title = self._title_from_url(final_url)
            return title, payload.decode('utf-8', errors='ignore'), final_url

        decoded = payload.decode('utf-8', errors='ignore')
        if '<html' in decoded.lower() or 'text/html' in content_type:
            extractor = ReadableHtmlExtractor()
            extractor.feed(decoded)
            extracted_title, text = extractor.extract()
            title = extracted_title or self._title_from_url(final_url)
            return title, text, final_url

        title = self._title_from_url(final_url)
        return title, decoded, final_url

    def _extract_remote_file(self, final_url: str, payload: bytes, suffix: str) -> tuple[str, str, str]:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
            temp_file.write(payload)
            temp_path = Path(temp_file.name)
        try:
            content = self._extract_text(temp_path, suffix)
        finally:
            temp_path.unlink(missing_ok=True)
        return self._title_from_url(final_url), content, final_url

    def _extract_text(self, path: Path, suffix: str) -> str:
        if suffix == '.pdf':
            reader = PdfReader(str(path))
            return '\n'.join((page.extract_text() or '') for page in reader.pages)
        if suffix == '.docx':
            doc = DocxDocument(str(path))
            return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        return path.read_text(encoding='utf-8', errors='ignore')

    def _chunk_text(self, content: str, target_words: int = 110) -> list[str]:
        paragraphs = [segment.strip() for segment in re.split(r'\n\s*\n', content) if segment.strip()]
        if not paragraphs:
            paragraphs = [content.strip()]
        chunks: list[str] = []
        current: list[str] = []
        current_words = 0
        for paragraph in paragraphs:
            words = paragraph.split()
            if current and current_words + len(words) > target_words:
                chunks.append('\n\n'.join(current))
                current = [paragraph]
                current_words = len(words)
            else:
                current.append(paragraph)
                current_words += len(words)
        if current:
            chunks.append('\n\n'.join(current))
        return chunks

    def _summarize(self, content: str, max_words: int = 24) -> str:
        words = content.strip().split()
        summary = ' '.join(words[:max_words])
        return summary + ('…' if len(words) > max_words else '')

    def _title_from_url(self, url: str) -> str:
        parsed = urlparse(url)
        segment = Path(parsed.path).stem.replace('-', ' ').replace('_', ' ').strip()
        if segment:
            return segment.title()
        return parsed.netloc.replace('www.', '')

    def _tokenize(self, query: str) -> list[str]:
        return [token for token in re.findall(r'[a-z0-9]+', query.lower()) if len(token) > 1]

    def _score_chunk(self, query_tokens: list[str], text: str) -> tuple[float, list[str]]:
        haystack = text.lower()
        matches: list[str] = []
        score = 0.0
        for token in query_tokens:
            if token in haystack:
                matches.append(token)
                score += 1.0 + (haystack.count(token) * 0.2)
        normalized = score / max(len(query_tokens), 1)
        return normalized, matches

    def _row_to_document(self, row) -> Document:
        return Document(
            id=row['id'],
            title=row['title'],
            type=row['type'],
            tags=database.loads(row['tags']) or [],
            date=row['date'],
            summary=row['summary'],
            content=row['content'],
            source_name=row.get('source_name'),
            source_kind=row.get('source_kind'),
            created_at=row['created_at'],
            updated_at=row['updated_at'],
        )

    def _row_to_chunk(self, row) -> Chunk:
        embedding = database.loads(row.get('embedding')) if row.get('embedding') else None
        return Chunk(
            id=row['id'],
            document_id=row['document_id'],
            document_title=row['document_title'],
            text=row['text'],
            position=row.get('position') or 0,
            metadata=database.loads(row['metadata']) or {},
            embedding=embedding,
        )


document_service = DocumentService()
